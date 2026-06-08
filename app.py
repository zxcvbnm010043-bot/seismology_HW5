# -*- coding: utf-8 -*-
import gradio as gr
import os
import glob
import numpy as np
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt

picture_dir = 'EE_picture'
image_paths = sorted(glob.glob(os.path.join(picture_dir, '*.jpg')))

def simulate_building_response(building_height, has_damper):
    t = np.linspace(0, 10, 500)
    if building_height == '高樓層 (Tall)':
        wn = 2.0 * np.pi / 2.0
        label_h = '高樓'
    else:
        wn = 2.0 * np.pi / 0.5
        label_h = '低樓'
    zeta = 0.15 if has_damper else 0.02
    wd = wn * np.sqrt(1 - zeta**2)
    response = np.exp(-zeta * wn * t) * np.cos(wd * t)
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=t, y=response, mode='lines', 
                             name=f'{label_h} + {"有" if has_damper else "無"}阻尼器',
                             line=dict(color='#FF4B4B' if has_damper else '#1C1C1C', width=3)))
    fig.update_layout(
        title=dict(text=f'建築物震動反應模擬 ({label_h}, {"有" if has_damper else "無"}阻尼器)', font=dict(size=18)),
        xaxis_title='時間 (秒)', yaxis_title='位移量 (相對能量)', template='plotly_white',
        margin=dict(l=40, r=40, t=60, b=40), height=400, yaxis=dict(range=[-1.1, 1.1])
    )
    return fig

process_text = """
### 1. 參觀過程
本次參觀的第一站是了解「國家地震工程研究中心」（NCREE）的核心使命。該中心致力於研發最先進的地震工程技術，從建築減震、強震觀測到災後防災評估，目標在於將天災對人類社會造成的傷亡與經濟損失減至最低。導覽員詳細介紹了台灣的地質背景，特別是位於歐亞板塊與菲律賓海板塊交界處的特殊位置，造就了頻繁的地震活動。

在基礎知識環節，我們討論了「彈性反彈理論」（Elastic Rebound Theory）。這個理論將地殼板塊比作一根巨大的彈簧，雖然每年僅移動數公分，但長時間的應力累積會在超過地殼強度時突然爆發釋放能量，造成強烈震動。隨後，導覽員透過生動的比喻與模型，解釋了「共振現象」對於建築物的威脅。例如，低評的地震波（長週期）容易與超高層建築產生共振，而高頻波（短週期）則對低矮透天厝威脅較大。

緊接著，我們深入探討了現代建築的三大「保命符」：
1. **調諧質量阻尼器 (TMD)**：類似台北 101 的金球，利用慣性來抵銷晃動；
2. **油壓阻尼器 (Viscous Damper)**：利用高黏滯性流體的阻尼係數來吸收動能，轉化為熱能消散；
3. **挫曲束制支撐 (BRB)**：傳統鋼構支撐在受拉與受壓時都能穩定吸收能量，外圍加上套管與塗層防止失效。

參觀的重頭戲是實體實驗室區。現場最令人震撼的是「多軸向測試系統」（MATS），該系統能模擬多維度的強烈位移與應力，測試各種新型隔震墊或建材在面臨 5500 公噸重壓下的表現。最後，前往七樓參觀模型，親自體驗當地動來襲時，傳統結構與先進隔震墊（如鉛芯橡膠支承）在穩定度上的天壤之別。
"""

reflection_text = """
### 3. 心得感想
站在 NCREE 展示的「土壤液化與地質風險潛勢圖」前，我心中感觸良多。特別是看見台北市士林區、社子島一帶被標示為高潛勢紅色區塊時，這讓我深刻體會到，雖然土地天生脆弱，但這正是工程技術必須介入的地方。如果我們無法改變居住的土地，那麼我們就必須強化居住的結構。

進度實驗區時那種「工業感」的嚴肅氣氛也讓我也印象深刻。每個人戴上工程安全帽的那一刻，彷彿參與了一場守衛城市安全的科學聖戰。看著那些冷冰冰的巨大鋼構與液壓機組，我卻感受到了一種守護生命的溫暖。尤其是 MATS 測試區那 5500 公噸的數據，這展現了人類在對抗極端天災時的不屈毅力。

此外，我也從中思考了一些更深層的工程實務問題。例如：當現今技術已經能做到全方位的隔震，使整棟建築與地盤分離時，地下的供水、供電與瓦斯管線要如何維持彈性連接而不斷裂？雖然相對於壯闊的地球，人類的大腦或許渺小如沙礫，但我們卻能透過科學方法，將原本混亂無序的天災，轉化為可以計算、分析並克服的挑戰。這場參觀，不只是關於冷冰冰的鋼筋混凝土，更是一場關於生存智慧的啟示錄。
"""

qa_text = """
### 4. 心得報告問題
**a. 101大樓的阻尼器主要任務是甚麼？為什麼 101 對大部分從遠方來的地震無感？**
台北 101 的調諧質量阻尼器（TMD）主要任務有二：一是「減震」，即在地震發生時吸收能量；二是「舒適」，減緩強風造成的高頻微幅晃動。
「無感」原因有三：
1. **能量衰減**：遠方波傳至台北盆地已大幅衰減；
2. **基盤穩定**：101 的基礎深入岩盤，地基穩固；
3. **動態消能**：即便產生微幅共振，強大的 TMD 也會迅速將能量轉化為阻尼器動能而非結構位移。

**b. 甚麼是建築物反應譜？甚麼是耐震設計？**
*   **建築物反應譜 (Response Spectrum)**：是一張結合了地震波特徵與建築動力特性的圖表。工程師可以藉此快速查詢特定高度的建築物在當地下地震時會承受多大的「地震力」。
*   **耐震設計 (Seismic Design)**：是一種包含結構配置與材料挑選的整體程序。核心理念為「小震不壞、中震可修、大震不倒」。

**c. 甚麼樣的建築物具有較高的 seismic risk ? 有哪些 NG 事項？**
*   **高 Seismic Risk 特徵**：位於活動斷層帶、土壤液化區；921 前建造的老舊建築；平面或立面形狀極不規則；擁有「軟弱底層」（如挑高店面）。
*   **影響抗震能力的 NG 事項**：
    1. **拆除承重牆/剪力牆**：隨意抽掉結構牆，破壞建築「脊椎」。
    2. **非法加蓋**：增加整棟樓的頂部質量，增大地震慣性力。
    3. **梁柱任意洗洞**：為了裝潢管線而切斷主鋼筋或縮減結構斷面。
    4. **未補強的老舊建築**：忽視結構裂縫或傾斜。
"""

def export_to_docx():
    output_docx = "earthquake_report.docx"
    doc = Document()
    doc.add_heading('國家地震工程研究中心 參觀心得報告', 0)
    
    doc.add_heading('1. 參觀過程', level=1)
    doc.add_paragraph(process_text.strip())
    
    doc.add_heading('2. 照片分享', level=1)
    if image_paths:
        try:
            doc.add_picture(image_paths[0], width=Inches(5))
            doc.add_paragraph("圖：參觀現場精華照片")
        except:
            pass
        
    doc.add_heading('3. 心得感想', level=1)
    doc.add_paragraph(reflection_text.strip())
    
    doc.add_heading('4. 問題解答', level=1)
    doc.add_paragraph(qa_text.strip())
    
    doc.save(output_docx)
    return output_docx

custom_css = '.container { max-width: 900px; margin: auto; } .section-box { border-left: 4px solid #FF4B4B; padding-left: 20px; margin-bottom: 40px; } .nav-btn { text-align: left !important; }'

with gr.Blocks(title='國家地震工程研究中心 專題報告') as demo:
    with gr.Row():
        with gr.Column(scale=1, min_width=200):
            gr.Markdown('## 📋 報告目錄')
            nav_process = gr.Button('1. 參觀過程', variant='secondary', elem_classes='nav-btn')
            nav_compare = gr.Button('2. 模擬比較', variant='secondary', elem_classes='nav-btn')
            nav_gallery = gr.Button('3. 照片分享', variant='secondary', elem_classes='nav-btn')
            nav_reflect = gr.Button('4. 心得感想', variant='secondary', elem_classes='nav-btn')
            nav_qa = gr.Button('5. 問題解答', variant='secondary', elem_classes='nav-btn')
            gr.Markdown("---")
            export_btn = gr.File(label="下載 Word 報告", visible=False)
            gen_btn = gr.Button("📄 生成並下載 Word 檔", variant="primary")
            
        with gr.Column(scale=4):
            gr.Markdown('# 🏢 國家地震工程研究中心 參觀心得報告')
            with gr.Group(visible=True) as sec_process:
                gr.Markdown(process_text, elem_classes='section-box')
            with gr.Group(visible=False) as sec_compare:
                gr.Markdown('### 📊 耐震模擬比較', elem_classes='section-box')
                gr.Markdown('調整下方參數，觀察不同建築高度與是否有阻尼器對於震動衰減的影響。')
                with gr.Row():
                    input_height = gr.Radio(['低樓層 (Short)', '高樓層 (Tall)'], label='建築物高度', value='高樓層 (Tall)')
                    input_damper = gr.Checkbox(label='安裝阻尼器 (Damper)', value=True)
                plot_output = gr.Plot()
                input_height.change(simulate_building_response, [input_height, input_damper], plot_output)
                input_damper.change(simulate_building_response, [input_height, input_damper], plot_output)
                demo.load(simulate_building_response, [input_height, input_damper], plot_output)
            with gr.Group(visible=False) as sec_gallery:
                gr.Markdown('### 📸 參觀照片集錦', elem_classes='section-box')
                if image_paths:
                    gr.Gallery(value=image_paths, label='現場照片', columns=[3], rows=[3], height='auto')
                else:
                    gr.Markdown('找不到圖片。')
            with gr.Group(visible=False) as sec_reflect:
                gr.Markdown(reflection_text, elem_classes='section-box')
            with gr.Group(visible=False) as sec_qa:
                gr.Markdown(qa_text, elem_classes='section-box')

    sections = [sec_process, sec_compare, sec_gallery, sec_reflect, sec_qa]
    def switch_section(idx):
        return [gr.update(visible=(i == idx)) for i in range(len(sections))]

    nav_process.click(fn=lambda: switch_section(0), outputs=sections)
    nav_compare.click(fn=lambda: switch_section(1), outputs=sections)
    nav_gallery.click(fn=lambda: switch_section(2), outputs=sections)
    nav_reflect.click(fn=lambda: switch_section(3), outputs=sections)
    nav_qa.click(fn=lambda: switch_section(4), outputs=sections)
    
    gen_btn.click(fn=export_to_docx, outputs=export_btn).then(
        fn=lambda: gr.update(visible=True), outputs=export_btn
    )

if __name__ == '__main__':
    print("Starting Gradio app...")
    demo.launch(ssr_mode=False)
