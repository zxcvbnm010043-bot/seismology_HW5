# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    print("Starting docx generation...")
    doc = Document()
    
    # 設定字體 (微軟正黑體 fallback)
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(12)
    
    # 標題
    title = doc.add_heading('國家地震工程研究中心 參觀心得報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 讀取文字內容
    content_file = 'EE_word.txt'
    if not os.path.exists(content_file):
        print("Error: EE_word.txt not found")
        return

    with open(content_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_section = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('1.參觀過程') or line.startswith('2.照片分享') or line.startswith('3.心得感想') or line.startswith('4.心得報告問題'):
            doc.add_heading(line, level=1)
            current_section = line
            
            # 如果是照片分享，插入幾張代表性照片
            if '2.照片分享' in line:
                pic_dir = 'EE_picture'
                if os.path.exists(pic_dir):
                    pics = [f for f in os.listdir(pic_dir) if f.lower().endswith(('.jpg', '.png'))]
                    if pics:
                        # 插入前三張作為代表
                        for p in pics[:3]:
                            try:
                                doc.add_picture(os.path.join(pic_dir, p), width=Inches(5))
                                doc.add_paragraph(f"圖：{p}", style='Caption')
                            except:
                                pass
        else:
            # 一般段落
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_name = '國家地震工程研究中心_參觀心得報告.docx'
    doc.save(output_name)
    print(f"Successfully generated: {output_name}")

if __name__ == "__main__":
    create_report()
